import os
import io
import sys
import glob
import json
import ctypes
import platform
import shutil
import datetime
import subprocess
import traceback
import threading
import urllib.request
import socket
import webbrowser
import customtkinter as ctk
import threading, queue, time
from subprocess import Popen, PIPE
from PIL import Image
from tkinter import filedialog
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Constants for Win32 API
GWL_STYLE = -16
WS_OVERLAPPEDWINDOW = 0x00CF0000
WS_POPUP = 0x80000000
SWP_FRAMECHANGED = 0x0020
SWP_NOMOVE = 0x0002
SWP_NOSIZE = 0x0001
SWP_NOZORDER = 0x0004

WM_NCLBUTTONDOWN = 0x00A1
HTCAPTION = 2

def resource_path(relative_path):
    try:
        if getattr(sys, 'frozen', False):
            return os.path.join(sys._MEIPASS, relative_path)
        else:
            return os.path.join(os.path.abspath("."), relative_path)
    except Exception as e:
        print(f"Error while accessing resource path: {e}")
        return None

def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = get_base_dir()
AUTOEXPORT_DIR = os.path.join(BASE_DIR, ".autoexport")
CONFIG_FILE = os.path.join(AUTOEXPORT_DIR, "cache.json")
VISUALS_DIR = os.path.join(AUTOEXPORT_DIR, "visuals")

THEMES_DIR = "assets/themes"

def get_theme_files():
    return sorted(glob.glob(os.path.join(THEMES_DIR, "*.json")))

def load_theme():
    try:
        with open(CONFIG_FILE, "r") as f:
            data = json.load(f)
            return data.get("theme")
    except Exception:
        return "blue"
    
def load_last_theme():
    try:
        with open(CONFIG_FILE, "r") as f:
            data = json.load(f)
            return data.get("last_theme")
    except Exception:
        return "blue"
    
def save_theme(theme_file):
    os.makedirs(os.path.dirname(CONFIG_FILE), exist_ok=True)
    try:
        with open(CONFIG_FILE, "r") as f:
            data = json.load(f)
    except Exception:
        data = {}
    data["theme"] = theme_file
    data["last_theme"] = load_theme()
    with open(CONFIG_FILE, "w") as f:
        json.dump(data, f, indent=4)

theme_window = None

def theme_selection_popup():
    global theme_window

    if theme_window is not None and theme_window.winfo_exists():
        theme_window.lift()
        theme_window.focus_force()
        return

    theme_files = get_theme_files()

    theme_window = ctk.CTkToplevel(app)
    theme_window.overrideredirect(True)
    #check_minimize()

    theme_window.attributes("-topmost", True)
    
    screen_width = theme_window.winfo_screenwidth()
    screen_height = theme_window.winfo_screenheight()

    button_count = len(theme_files)

    button_width = 240
    button_height = 30
    padding = 2

    window_height = button_count * (button_height + padding) + padding + 120
    window_width = button_width + padding * 2

    close_button = ctk.CTkButton(
        theme_window,
        width=36,
        height=30,
        text="",
        image=close_icon,
        hover_color="#AA0000",
        command=theme_window.destroy
    )
    close_button.place(x=int(window_width / 2 - 18), y=window_height - 35)

    current_theme_path = load_theme()
    current_theme_name = os.path.basename(current_theme_path).replace(".json", "") if current_theme_path else "Default/Blue"

    current_label = ctk.CTkLabel(
        theme_window,
        text=f"Current Theme: {current_theme_name}",
        font=("Segoe UI", 14, "bold"),
        text_color="#BBBBBB"
    )
    current_label.pack(pady=(10, 5))

    theme_window.geometry(f"{window_width}x{window_height}+{(screen_width - window_width) // 2}+{(screen_height - window_height) // 2}")

    def check_minimize(_=None):
        try:
            if 'theme_window' in globals() and isinstance(theme_window, ctk.CTkToplevel):
                if theme_window.winfo_exists():
                    if app.state() == 'iconic':
                        theme_window.withdraw()
                    else:
                        theme_window.deiconify()
                    app.after(300, check_minimize) 
                    return
        except Exception as e:
            print(f"[check_minimize] Ignored error: {e}")

    app.bind("<Unmap>", check_minimize)  
    app.bind("<Map>", check_minimize)    

    def on_select_theme(selected_theme):
        theme_files = get_theme_files()
        theme_file_path = None
        for theme_file in theme_files:
            if os.path.basename(theme_file).replace(".json", "") == selected_theme:
                theme_file_path = theme_file
                break

        if theme_file_path:
            save_theme(theme_file_path)
            theme_window.destroy()
            show_restart_prompt()

    def restart_app():
        if getattr(sys, 'frozen', False):
            executable = sys.executable
            # Delay execution to allow current app to exit cleanly
            subprocess.Popen([executable], close_fds=True)
        else:
            python = sys.executable
            script = os.path.abspath(sys.argv[0])
            subprocess.Popen([python, script, *sys.argv[1:]], close_fds=True)

        # Exit after a short delay to avoid clashing with subprocess launch
        threading.Thread(target=lambda: (time.sleep(0.5), os._exit(0))).start()

    
    def show_restart_prompt():
        popup = ctk.CTkToplevel()
        popup.title("Restart Required")
        popup.geometry("300x110")
        popup.resizable(False, False)
        popup.attributes("-topmost", True)
        popup.overrideredirect(True)
        
        label = ctk.CTkLabel(
            popup,
            text="Restart is required to apply the theme.\nRestart now?",
        )
        label.pack(pady=(20, 10))

        btn_frame = ctk.CTkFrame(popup, fg_color="transparent", border_width=0)
        btn_frame.pack(pady=10)

        def on_yes():
            popup.destroy()
            app.destroy()
            restart_app()

        def on_no():
            popup.destroy()

        yes_btn = ctk.CTkButton(btn_frame, text="Yes", command=on_yes)
        yes_btn.pack(side="left", padx=10)

        no_btn = ctk.CTkButton(btn_frame, text="No", command=on_no)
        no_btn.pack(side="right", padx=10)

        popup.update_idletasks()
        x = app.winfo_x() + (app.winfo_width() - popup.winfo_width()) // 2
        y = app.winfo_y() + (app.winfo_height() - popup.winfo_height()) // 2
        popup.geometry(f"+{x}+{y}")


    def get_theme_properties(theme_file):
        try:
            with open(theme_file, "r") as f:
                theme_data = json.load(f)
                button_properties = theme_data.get("CTkButton", {})
                font_properties = theme_data.get("CTkFont", {})
                return button_properties, font_properties
        except Exception:
            return {}, {}

    def get_font_for_platform(font_properties):
        current_platform = platform.system()
        if current_platform == "Darwin":  # macOS
            return font_properties.get("macOS", {})
        elif current_platform == "Windows":  # Windows
            return font_properties.get("Windows", {})
        else:  # Linux
            return font_properties.get("Linux", {})

    if not theme_files:
        no_theme_label = ctk.CTkLabel(
            theme_window,
            text=f"No theme files found.\nRead release notes.",
            font=("Segoe UI", 12, "normal"),
            text_color="#BBBBBB"
        )
        no_theme_label.pack()
    else:
        for theme_file in theme_files:
            theme_name = os.path.basename(theme_file).replace(".json", "")
            button_properties, font_properties = get_theme_properties(theme_file)  

            corner_radius = button_properties.get("corner_radius", 6) 
            border_width = button_properties.get("border_width", 0)  
            fg_color = button_properties.get("fg_color", ["#3B8ED0", "#1F6AA5"])  
            hover_color = button_properties.get("hover_color", ["#36719F", "#144870"]) 
            border_color = button_properties.get("border_color", ["#3E454A", "#949A9F"])  
            text_color = button_properties.get("text_color", ["#DCE4EE", "#DCE4EE"]) 
            text_color_disabled = button_properties.get("text_color_disabled", ["gray74", "gray60"]) 

            font_info = get_font_for_platform(font_properties)
            font_family = font_info.get("family", "Roboto")  
            font_size = font_info.get("size", 13)  
            font_weight = font_info.get("weight", "normal")  

            is_current = (theme_name == current_theme_name)
            display_name = f"✓ {theme_name}" if is_current else theme_name

            theme_button = ctk.CTkButton(
                theme_window,
                text=display_name,
                width=button_width,
                height=button_height,
                corner_radius=corner_radius,
                border_width=border_width,
                fg_color="#888888" if is_current else fg_color,
                hover_color="#666666" if is_current else hover_color,
                border_color=border_color,
                text_color=text_color,
                text_color_disabled="white",
                font=(font_family, font_size, font_weight),
                state="disabled" if is_current else "normal",
                command=lambda theme=theme_name: on_select_theme(theme)
            )

            theme_button.pack(pady=padding)
        
def get_compiled_path(language_ext: str, filename: str):
    subdir = os.path.join(AUTOEXPORT_DIR, f".{language_ext}")
    os.makedirs(subdir, exist_ok=True)
    return os.path.join(subdir, filename)

def ensure_export_dirs():
    os.makedirs(AUTOEXPORT_DIR, exist_ok=True)
    os.makedirs(VISUALS_DIR, exist_ok=True)

def is_internet_available():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=2)
        return True
    except OSError:
        return False


def check_for_update(local_version="2.3"):
    if not is_internet_available():
        return None

    try:
        url = "https://raw.githubusercontent.com/Shresth-Dwivedi/AutoExport/main/version.txt"
        with urllib.request.urlopen(url) as response:
            latest_version = response.read().decode("utf-8").strip()

        if latest_version != local_version:
            return f"Update available: v{latest_version}"
        return None
    except Exception:
        return None

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {}

def append_to_json(key, value, path=".autoexport/cache.json"):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        with open(path, "r") as f: data = json.load(f)
    except: data = {}
    data[key] = value
    with open(path, "w") as f: json.dump(data, f, indent=4)

def get_sorted_file_paths(paths, sort_by):
        if sort_by == "Date Created":
            return sorted(paths, key=lambda p: os.path.getctime(p))
        elif sort_by == "File Size":
            return sorted(paths, key=lambda p: os.path.getsize(p))
        elif sort_by == "Language":
            return sorted(paths, key=lambda p: os.path.splitext(p)[1])
        return paths

def export_programs(author, selected_files, output_path, metadata=None, execution_outputs=None, visual_outputs=None):
    if not selected_files:
        return "No files selected for export."

    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.lift()
    progress_bar.set(0)
    progress_bar.update()
    status_label.configure(text="Starting export...")
    status_label.lift()
    try:
        if app.winfo_exists():
            app.update_idletasks()
    except:
        pass

    doc = Document()

    if 'CodeBlock' not in doc.styles:
        style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.right_indent = Inches(0.3)

    doc.add_heading("Programming Assignment", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Author: {author}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Exported on: {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if metadata:
        doc.add_paragraph(f"Instructor: {metadata.get('instructor')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Course: {metadata.get('course')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Semester: {metadata.get('semester')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if metadata.get("notes"):
            doc.add_paragraph(f"Notes: {metadata.get('notes')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()
    doc.add_paragraph("Table of Contents", style="Heading 1")
    doc.add_paragraph("→ Update fields in Word via References → Table of Contents → Update Field.")
    doc.add_page_break()

    # Apply sorting before export
    selected_sort = sort_selector.get()
    selected_files = get_sorted_file_paths(selected_files, selected_sort)

    total_files = len(selected_files)

    for i, path in enumerate(selected_files, 1):
        try:
            filename = os.path.basename(path)
            status_label.configure(text=f"Exporting...")
            app.update_idletasks()

            with open(path, "r", encoding="utf-8") as f:
                content = f.read()

            modified_time = datetime.datetime.fromtimestamp(os.path.getmtime(path)).strftime("%d %b %Y, %H:%M")

            doc.add_heading(filename, level=2)
            doc.add_paragraph(f"Last Modified: {modified_time}", style="Normal")

            for line in content.splitlines():
                doc.add_paragraph(line, style='CodeBlock')

            #status_label.configure(text=f"Writing output for: {filename}")
            #app.update_idletasks()

            doc.add_paragraph("Output:", style="Heading 2")
            output = (
                execution_outputs.get(path)
                or execution_outputs.get(filename)
                or "Output not available."
            )

            for line in output.splitlines():
                doc.add_paragraph(line, style='CodeBlock')

            if visual_outputs:
                images = visual_outputs.get(path, []) or visual_outputs.get(filename, [])
                if images:
                    #status_label.configure(text=f"Adding visuals for: {filename}")
                    #app.update_idletasks()

                    doc.add_paragraph("Visual Output(s):", style="Heading 2")
                    for image_path in images:
                        if os.path.exists(image_path):
                            try:
                                doc.add_picture(image_path, width=Inches(5.5))
                                doc.add_paragraph(f"[{os.path.basename(image_path)}]", style='Normal')
                            except Exception as e:
                                doc.add_paragraph(f"[Failed to load image: {image_path} — {e}]", style='Normal')
                        else:
                            doc.add_paragraph(f"[Missing image file: {image_path}]", style='Normal')

        except Exception as e:
            print(f"Skipped file {path}: {e}")
            doc.add_paragraph(f"[Failed to process file: {path}]", style="Normal")

        doc.add_page_break()

        progress = i / total_files
        progress_bar.set(progress)
        try:
            if app.winfo_exists():
                app.update_idletasks()
        except:
            pass

    status_label.configure(text="Saving document...")
    app.update_idletasks()

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Failed to save Word file: {e}")
        if app.winfo_exists():
            status_label.configure(text="")
            app.after(1000, lambda: progress_bar.place_forget())
            return False 

    try:
        status_label.configure(text="Opening exported file...")
        app.update_idletasks()
        subprocess.Popen(["start", "", output_path], shell=True)
    except Exception as e:
        print(f"Could not open Word file automatically: {e}")

    if app.winfo_exists():
        status_label.configure(text="")
        app.after(1000, lambda: progress_bar.place_forget())
    
    app.update_idletasks

    return True 

def terminal_input_session(program_name, input_prompts):
    inputs = []

    dialog = ctk.CTkToplevel()
    dialog.title(f"Terminal Input – {program_name}")
    dialog.geometry("600x400")
    dialog.resizable(False, False)
    
    dialog.update_idletasks()
    w, h = 600, 400
    x = (dialog.winfo_screenwidth() // 2) - (w // 2)
    y = (dialog.winfo_screenheight() // 2) - (h // 2)
    dialog.geometry(f"{w}x{h}+{x}+{y}")
    dialog.grab_set()

    output_box = ctk.CTkTextbox(dialog, height=280, wrap="word", font=("Consolas", 12))
    output_box.pack(fill="both", expand=True, padx=10, pady=(10, 5))
    output_box.configure(state="disabled")

    entry_var = ctk.StringVar()
    entry = ctk.CTkEntry(dialog, textvariable=entry_var, font=("Consolas", 12))
    entry.pack(fill="x", padx=10, pady=(0, 10))
    entry.focus_set()

    current_index = [0]

    def print_to_terminal(text):
        output_box.configure(state="normal")
        output_box.insert("end", text + "\n")
        output_box.configure(state="disabled")
        output_box.see("end")

    def on_submit(*_):
        if current_index[0] < len(input_prompts):
            answer = entry_var.get()
            inputs.append(answer)

            print_to_terminal(f"{input_prompts[current_index[0]]}{answer}")
            entry_var.set("")
            current_index[0] += 1

            if current_index[0] < len(input_prompts):
                print_to_terminal(input_prompts[current_index[0]])
            else:
                dialog.destroy()

    entry.bind("<Return>", on_submit)

    print_to_terminal(f"Running: {program_name}")
    print_to_terminal(input_prompts[0])

    dialog.wait_window()
    return inputs

def detect_language_and_command(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    filedir = os.path.dirname(file_path)
    basename = os.path.splitext(filename)[0]

    if ext == ".py" and shutil.which("python"):
        return "Python", ["python", file_path], None

    elif ext == ".cpp" and shutil.which("g++"):
        output = os.path.join(filedir, f"{basename}.exe" if os.name == "nt" else basename)
        return "C++", ["g++", file_path, "-o", output], output

    elif ext == ".c" and shutil.which("gcc"):
        output = os.path.join(filedir, f"{basename}.exe" if os.name == "nt" else basename)
        return "C", ["gcc", file_path, "-o", output], output

    elif ext == ".java" and shutil.which("javac") and shutil.which("java"):
        return "Java", ["javac", file_path], basename

    elif ext == ".js" and shutil.which("node"):
        return "JavaScript", ["node", file_path], None

    elif ext == ".go" and shutil.which("go"):
        return "Go", ["go", "run", file_path], None

    elif ext == ".rb" and shutil.which("ruby"):
        return "Ruby", ["ruby", file_path], None

    elif ext == ".php" and shutil.which("php"):
        return "PHP", ["php", file_path], None

    elif ext == ".cs" and shutil.which("csc"):
        exe_name = os.path.join(filedir, f"{basename}.exe")
        return "C#", ["csc", "/out:" + exe_name, file_path], exe_name

    return None, [], None

def create_terminal_window(title="Terminal"):

    dialog = ctk.CTkToplevel(master=app)
    dialog.geometry("650x400")
    dialog.resizable(False, False)
    dialog.overrideredirect(1)

    dialog.update_idletasks()
    w, h = 650, 400
    x = (dialog.winfo_screenwidth() // 2) - (w // 2)
    y = (dialog.winfo_screenheight() // 2) - (h // 2)
    dialog.geometry(f"{w}x{h}+{x}+{y}")
    terminal_closed = [False] 
    dialog.grab_set()

    output_box = ctk.CTkTextbox(dialog, font=("Consolas", 12), wrap="word")
    output_box.pack(fill="both", expand=True, padx=10, pady=(10, 5))
    output_box.configure(state="disabled")

    input_var = ctk.StringVar()
    input_entry = ctk.CTkEntry(dialog, textvariable=input_var, font=("Consolas", 12))
    input_entry.pack(fill="x", padx=10, pady=(0, 10))

    prompt_event = threading.Event()
    input_queue = queue.Queue()
    collected_output = []  
    
    def write(text):
        output_box.configure(state="normal")
        output_box.insert("end", text)
        output_box.configure(state="disabled")
        output_box.see("end")
        collected_output.append(text)   

    def wait_for_prompt(tag="input"):
        def enable_input():
            prompt_event.set()
            input_entry.focus_set()

        dialog.after(100, enable_input)
        
    def get_next_input():
        prompt_event.wait(100)  
        value = input_queue.get()  
        prompt_event.clear()
        return value

    def on_submit(event=None):
        # if not prompt_event.is_set():
        #     return
        value = input_var.get()
        input_var.set("")
        input_queue.put(value)
        write(value + "\n")  
        
    input_entry.bind("<Return>", on_submit)

    def destroy():
        terminal_closed[0] = True
        try:
            dialog.destroy()
        except:
            pass

    def is_closed():
        return terminal_closed[0]

    def exists():
        return bool(dialog.winfo_exists())

    return type("Terminal", (), {
        "write": staticmethod(write),
        "wait_for_prompt": staticmethod(wait_for_prompt),
        "get_next_input": staticmethod(get_next_input),
        "destroy": staticmethod(destroy),
        "is_closed": staticmethod(is_closed),
        "exists": staticmethod(exists),
        "collected_output": collected_output,
    })()


def execute_with_input_detection(cmd_list, program_name, force_terminal=False, timeout=60):
    
    q = queue.Queue()
    collected_output = []
    buffer = ""
    terminal = None
    show_terminal = False

    if force_terminal:
        terminal = create_terminal_window(program_name)
        show_terminal = True

    process = Popen(cmd_list, stdin=PIPE, stdout=PIPE, stderr=PIPE, text=True, bufsize=0,  creationflags=subprocess.CREATE_NO_WINDOW)

    def read_stream(stream, tag):
        for chunk in iter(lambda: stream.read(1), ''):
            q.put((tag, chunk))
        q.put((tag, None))

    threading.Thread(target=read_stream, args=(process.stdout, "stdout"), daemon=True).start()
    threading.Thread(target=read_stream, args=(process.stderr, "stderr"), daemon=True).start()

    last_output_time = time.time()
    silence_threshold = 1.5
    show_terminal = force_terminal
    start_time = time.time()
    input_waiting = False
    
    import random
    roasts = [
        "Agle surgical strike ki script bann rahi hai kya?",
        "I've seen glaciers move faster.",
        "Rome fell faster than this executed.",
        "This execution speed is sponsored by Indian judiciary.",
        "Modiji finished another Mann Ki Baat, and this still runs!",
        "Bhaiya, kya ye code acche dino ka wait kar raha hai?",
        "Dekh raha hai Binod, kitna slow chal raha hai code.",
        "Beta, aap chor do coding karna. PLEASE...",
        "Sorry for the wait. Aapka yeh program hi jhaatu hai.",
        "Abey jaldi run hoja, kal subah panvel nikalna hai!",
        "Is this code powered by snails?"
    ]
    roast_shown = False

    while True:
        if (time.time() - start_time > 35) and not show_terminal and not roast_shown:
            roast_shown = True
            status_label.lift()
            status_label.configure(text=random.choice(roasts))
            app.update_idletasks()

        if (time.time() - start_time > timeout) and not show_terminal:
            try:
                process.terminate()
                status_label.configure(text="")
                return "[⚠] Program timed out. Possible infinite loop or took very long time to execute.\n"
            except Exception:
                status_label.configure(text="")
                return "[⚠] Failed to terminate after timeout.\n"
            
        try:
            while not q.empty():
                tag, chunk = q.get_nowait()
                if chunk is None:
                    continue

                collected_output.append(chunk)
                buffer += chunk

                if show_terminal:
                    terminal.write(chunk if tag == "stdout" else f"[stderr] {chunk}")
                
                last_output_time = time.time()
                input_waiting = False

            if show_terminal and not input_waiting and (time.time() - last_output_time > silence_threshold):
                input_waiting = True
                terminal.wait_for_prompt("input")
                user_input = terminal.get_next_input()
                if user_input is None:
                    break
                process.stdin.write(user_input + "\n")
                process.stdin.flush()

                collected_output.append(user_input + "\n")

                last_output_time = time.time()
                input_waiting = False

            if process.poll() is not None:
                break

        except Exception:
            break

    process.wait()

    if show_terminal and terminal:
        terminal.write("\n[✔] Program finished. Press Enter to close...\n")
        terminal.wait_for_prompt("exit")
        _ = terminal.get_next_input()
        terminal.destroy()

    return "".join(collected_output)

def program_requires_input(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read().lower()
            ext = os.path.splitext(file_path)[1].lower()

            input_indicators = []

            if ext == ".py":
                input_indicators = [
                    "input(", "sys.stdin.read", "sys.stdin.readline", "raw_input(",
                ]

            elif ext == ".java":
                input_indicators = [
                    "scanner", ".next", "system.in", "bufferedreader", ".readline",
                    "datainputstream", ".read(", ".readline("
                ]

            elif ext in [".c", ".cpp"]:
                input_indicators = [
                    "scanf", "getchar", "fgets", "gets(", "cin", "std::cin", "stdin",
                    "read(", "fgetc(", "getch("
                ]

            elif ext == ".cs":
                input_indicators = [
                    "console.readline", "streamreader", "textreader", ".readline", ".read("
                ]

            elif ext == ".js":
                input_indicators = [
                    "prompt(", "readline(", "process.stdin", "fs.createinterfacestream", "require('readline')"
                ]

            elif ext == ".go":
                input_indicators = [
                    "fmt.scan", "fmt.scanln", "bufio", "os.stdin", ".readstring"
                ]

            elif ext == ".rb":
                input_indicators = [
                    "gets", "stdin.gets", "readline"
                ]

            elif ext == ".php":
                input_indicators = [
                    "fgets(stdin", "readline(", "stream_socket_client(\"php://stdin", "$stdin"
                ]

            return any(keyword in code for keyword in input_indicators)

    except Exception:
        return False

            
def run_files_and_capture_output(file_paths):
    output_log = {}
    visual_outputs = {}

    os.makedirs(VISUALS_DIR, exist_ok=True)

    total_files = len(file_paths)

    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.lift()
    progress_bar.set(0.0)
    app.update()

    for i, path in enumerate(file_paths):
        buffer = io.StringIO()
        filename = os.path.basename(path)
        visuals_for_file = []

        status_label.lift()
        status_label.configure(text=f"Processing {filename}")
        app.update_idletasks()

        lang, compile_cmd, runtime_targe = detect_language_and_command(path)

        if not lang:
            output_log[path] = "⚠️ Unsupported file type or required toolchain not found."
            visual_outputs[path] = []
            continue

        try:
            try:
                import matplotlib
                matplotlib.use("Agg")
                import matplotlib.pyplot as plt

                original_show = plt.show
                original_savefig = plt.savefig

                def patched_show(*args, **kwargs):
                    nonlocal visuals_for_file
                    fname = f"{filename}_auto_figure_{len(visuals_for_file)+1}.png"
                    save_path = os.path.join(VISUALS_DIR, fname)
                    visuals_for_file.append(save_path)
                    original_savefig(save_path, *args, **kwargs)
                    plt.clf()
                    plt.cla()

                def patched_savefig(fname=None, *args, **kwargs):
                    if not fname:
                        fname = f"{filename}_figure_{len(visuals_for_file)+1}.png"
                    else:
                        fname = os.path.basename(fname)
                    save_path = os.path.join(VISUALS_DIR, fname)
                    visuals_for_file.append(save_path)
                    original_savefig(save_path, *args, **kwargs)

                plt.show = patched_show
                plt.savefig = patched_savefig
            except ImportError:
                plt = None

            requires_input = program_requires_input(path)
            ext = os.path.splitext(path)[1].lower()
            basename = os.path.splitext(os.path.basename(path))[0]

            if lang == "Python":
                status_label.configure(text=f"Running {filename}")
                app.update_idletasks()
                if requires_input:
                    output = execute_with_input_detection(["python", path], filename, force_terminal=True)
                    buffer.write(output)
                else:
                    global_namespace = {"__name__": "__main__"}
                    with open(path, "r", encoding="utf-8") as f:
                        code = f.read()
                    with redirect_stdout(buffer):
                        exec(compile(code, path, 'exec'), global_namespace)

            elif lang == "Java":
                status_label.configure(text=f"⚙ Compiling {filename}")
                app.update_idletasks()
                output_dir = os.path.dirname(get_compiled_path("java", "dummy.class"))
                compile_cmd = ["javac", "-d", output_dir, path]
                compile_proc = subprocess.run(compile_cmd, capture_output=True, text=True, creationflags=subprocess.CREATE_NO_WINDOW)
                if compile_proc.stderr:
                    buffer.write("Compile error:\n" + compile_proc.stderr)
                else:
                    status_label.configure(text=f"Running {filename}")
                    app.update_idletasks()
                    main_class = basename
                    output = execute_with_input_detection(["java", "-cp", output_dir, main_class], filename, force_terminal=requires_input)
                    buffer.write(output)

            elif lang in ["C", "C++", "C#", "Go", "Rust"]:
                status_label.configure(text=f"⚙ Compiling {filename}")
                app.update_idletasks()
                exe_name = basename + (".exe" if os.name == "nt" else "")
                output_path = get_compiled_path(ext[1:], exe_name)

                if ext == ".c":
                    compile_cmd = ["gcc", path, "-o", output_path]
                elif ext == ".cpp":
                    compile_cmd = ["g++", path, "-o", output_path]
                elif ext == ".cs":
                    compile_cmd = ["csc", "/out:" + output_path, path]
                elif ext == ".go":
                    compile_cmd = ["go", "build", "-o", output_path, path]
                elif ext == ".rs":
                    compile_cmd = ["rustc", path, "-o", output_path]

                subprocess.run(compile_cmd, capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)
                if os.path.exists(output_path):
                    status_label.configure(text=f"Running {filename}")
                    app.update_idletasks()
                    output = execute_with_input_detection([output_path], filename, force_terminal=requires_input)
                    buffer.write(output)
                else:
                    buffer.write("Compiled binary not found.\n")

            elif lang in ["JavaScript", "Ruby", "PHP"]:
                status_label.configure(text=f"Running {filename}")
                app.update_idletasks()
                output = execute_with_input_detection(compile_cmd, filename, force_terminal=requires_input)
                buffer.write(output)

            else:
                buffer.write("Unknown language execution path.\n")

        except Exception:
            buffer.write("Error during execution:\n")
            buffer.write(traceback.format_exc())

        if plt:
            plt.show = original_show
            plt.savefig = original_savefig

        output_log[path] = buffer.getvalue().strip()
        visual_outputs[path] = visuals_for_file

        progress = (i + 1) / total_files
        progress_bar.set(progress)
        app.update_idletasks()

    return output_log, visual_outputs

def attach_tooltip(widget, text, delay=1500):
    tooltip = None
    after_id = None

    try:
        fg_color_hex = json.load(open(load_theme()))["CTkButton"]["fg_color"][0]
        btn_color_hex = json.load(open(load_theme()))["CTkButton"]["text_color"][0]
    except Exception:
        fg_color_hex = "#3B8ED0"
        btn_color_hex = "#DCE4EE"

    def show_tooltip():
        nonlocal tooltip
        if not tooltip:
            tooltip = ctk.CTkLabel(widget.master, text=text, text_color=btn_color_hex, font=("Segoe UI", 10),corner_radius=8, fg_color=fg_color_hex)
            x = widget.winfo_x() - 35
            y = widget.winfo_y() - 18
            tooltip.place(x=x, y=y)

    def on_enter(event):
        nonlocal after_id
        after_id = widget.after(delay, show_tooltip)

    def on_leave(event):
        nonlocal tooltip, after_id
        if after_id:
            widget.after_cancel(after_id)
            after_id = None
        if tooltip:
            tooltip.destroy()
            tooltip = None

    widget.bind("<Enter>", on_enter)
    widget.bind("<Leave>", on_leave)


def bind_validation(entry, field_name):
    def validate(_event):
        value = entry.get().strip()

        def is_valid_author(name):
            return len(name) >= 4 and all(c.isalpha() or c.isspace() for c in name)

        def is_valid_output_file(name):
            return name.lower().endswith(".docx") and len(name) > 5

        def is_valid_folder_paths(paths):
            return any(os.path.isdir(path.strip()) for path in paths.split(";"))

        def is_valid_output_folder(folder):
            return folder == "" or os.path.isdir(folder)
    
        if field_name == "author":
            is_valid = is_valid_author(value)
        elif field_name == "prog_folder":
            is_valid = is_valid_folder_paths(value)
        elif field_name == "output_folder":
            is_valid = is_valid_output_folder(value)
        elif field_name == "output_file":
            is_valid = is_valid_output_file(value)
        else:
            is_valid = bool(value)  

        highlight_required(entry, error=not is_valid)
    entry.bind("<KeyRelease>", validate)


def highlight_required(entry, error: bool = True):
    border_color = "red" if error else "green"
    entry.configure(border_color=border_color, border_width=2)

def show_confetti(success=True, message="EXPORTED SUCCESSFULLY"):
    color = "green" if success else "red"
    #emoji = "🎉" if success else "❌"
    confetti = ctk.CTkLabel(details_frame, text=message, font=("Segoe UI", 18), text_color=color)
    confetti.place(relx=0.5, rely=0.95, anchor="center")
    app.after(3000, confetti.destroy)

def load_mode(path=CONFIG_FILE):
    try:
        with open(path, "r") as f:
            data = json.load(f)
            return data.get("mode", "dark")
    except:
        return "dark" 
    
class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        #icon for the app
        icon_path = resource_path(os.path.join("assets/icons/", "app.ico"))
        self.after(1000,self.iconbitmap(default=icon_path))

        window_width = 850
        window_height = 680

        # Screen dimensions
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        vertical_offset = 40

        x = (screen_width // 2) - (window_width // 2)
        y = max(0, (screen_height // 2) - (window_height // 2) - vertical_offset)

        self.geometry(f"{window_width}x{window_height}+{x}+{y}")
        self.resizable(False, False)
        self.make_window_frameless()

        self.bind("<ButtonPress-1>", self.start_move)

    def start_move(self, event):
        hwnd = ctypes.windll.user32.GetParent(self.winfo_id())
        ctypes.windll.user32.ReleaseCapture()
        ctypes.windll.user32.PostMessageW(hwnd, WM_NCLBUTTONDOWN, HTCAPTION, 0)

    def make_window_frameless(self):
        hwnd = ctypes.windll.user32.GetParent(self.winfo_id())
        style = ctypes.windll.user32.GetWindowLongW(hwnd, GWL_STYLE)
        style = style & ~WS_OVERLAPPEDWINDOW | WS_POPUP
        ctypes.windll.user32.SetWindowLongW(hwnd, GWL_STYLE, style)
        ctypes.windll.user32.SetWindowPos(hwnd, 0, 0, 0, 0,
                                          SWP_NOMOVE | SWP_NOSIZE | SWP_NOZORDER | SWP_FRAMECHANGED)

def show_theme_error_popup():
    popup = ctk.CTkToplevel()
    popup.title("Restart Required")
    popup.geometry("300x110")
    popup.resizable(False, False)
    popup.attributes("-topmost", True)
    popup.overrideredirect(True)
    
    label = ctk.CTkLabel(
        popup,
        text="Theme file corrupted.\nCannot load theme!",
    )
    
    label.pack(pady=(20, 10))

    btn_frame = ctk.CTkFrame(popup, fg_color="transparent", border_width=0)
    btn_frame.pack(pady=10)
    
    def on_ok():
        popup.destroy()
    
    ok_btn = ctk.CTkButton(btn_frame, text="Ok", command=on_ok)
    ok_btn.pack()

    popup.update_idletasks()
    x = app.winfo_x() + (app.winfo_width() - popup.winfo_width()) // 2
    y = app.winfo_y() + (app.winfo_height() - popup.winfo_height()) // 2
    popup.geometry(f"+{x}+{y}")

def modern_gui():
    global app, progress_bar, status_label, sort_selector, close_icon, details_frame, file_selection_state
    file_selection_state = {}
    config = load_config()

    initial_theme = load_theme()
    last_theme = load_last_theme()

    try:
        if initial_theme and os.path.exists(initial_theme):
            ctk.set_default_color_theme(initial_theme)
    except Exception:
        if last_theme:
            ctk.set_default_color_theme(last_theme)
            save_theme(last_theme)

    ctk.set_appearance_mode(load_mode())

    app = App()
    app.title("AutoExport")
    app.lift()
    app.focus_set()
    app.attributes('-topmost', True)
    app.after(1, lambda: app.attributes('-topmost', False))
    
    try:
        with open(CONFIG_FILE, "r") as f:
            config = json.load(f)
            theme = config.get("theme", "blue")
            last_theme = config.get("last_theme", "blue")
    except Exception:
        theme = last_theme = "blue"

    if last_theme == initial_theme:
        corrupted = False
        if theme != "blue" and not os.path.exists(os.path.join("themes", theme + ".json")):
            corrupted = True
        if last_theme != "blue" and not os.path.exists(os.path.join("themes", last_theme + ".json")):
            corrupted = True
        if corrupted:
            show_theme_error_popup()

    # Icons for labels and buttons
    close_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/close.png")), size=(20, 20))
    minimize_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/minimize.png")), size=(20, 20))
    
    theme_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/theme.png")), size=(20, 20))
    light_theme_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/light.png")), size=(20, 20))
    dark_theme_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/dark.png")), size=(20, 20))
    bell_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/bell.png")), size=(20, 20))
    clipboard_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/clip.png")), size=(20, 20))
    checkmark_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/check.png")), size=(20, 20))
    uncheck_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/uncheck.png")), size=(20, 20))
    back_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/back.png")), size=(20, 20))
    folder_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/open-folder.png")), size=(20, 20))

    woof_img = ctk.CTkImage(Image.open(resource_path("assets/buddy/woof.png")), size=(64*2, 64))

    # Convert .svg to .png
    import cairosvg
    import xml.etree.ElementTree as ET

    def load_svg_icon(svg_path, color_hex, size=32):
        with open(svg_path, 'r', encoding='utf-8') as f:
            svg_content = f.read()

        # Parse SVG XML
        root = ET.fromstring(svg_content)

        for path in root.iter('{http://www.w3.org/2000/svg}path'):
            path.set('fill', color_hex)

        svg_colored = ET.tostring(root, encoding='unicode')

        png_bytes = cairosvg.svg2png(bytestring=svg_colored, output_width=size, output_height=size)
        image = Image.open(io.BytesIO(png_bytes))

        return ctk.CTkImage(dark_image=image, light_image=image, size=(size, size))
    
    try:
        color_hex = json.load(open(load_theme()))["CTkButton"]["fg_color"][0]
    except Exception:
        color_hex = "#3B8ED0"

    # Social icons
    github_dark_icon = load_svg_icon(resource_path("assets/icons/github.svg"), color_hex)
    linkedin_dark_icon = load_svg_icon(resource_path("assets/icons/linkedin.svg"), color_hex)
    X_dark_icon = load_svg_icon(resource_path("assets/icons/X.svg"), color_hex)
    bluesky_dark_icon = load_svg_icon(resource_path("assets/icons/bluesky.svg"), color_hex)
    youtube_dark_icon = load_svg_icon(resource_path("assets/icons/youtube.svg"), color_hex)

    github_light_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/gitHub-light.png")), size=(32, 32))
    linkedin_light_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/linkedin-light.png")), size=(32, 32))
    X_light_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/X-light.png")), size=(20, 20))
    bluesky_light_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/bluesky-light.png")), size=(32, 32))
    youtube_light_icon = ctk.CTkImage(Image.open(resource_path("assets/icons/youtube-light.png")), size=(32, 32))

    close_button = ctk.CTkButton(
        app,
        width=36,
        height=30,
        text="",
        image=close_icon,
        hover_color="#AA0000",
        command=app.destroy
    )
    close_button.place(relx=0.05, rely=0.025, anchor="nw")

    minimize_button = ctk.CTkButton(
        app,
        width=36,
        height=30,
        text="",
        image=minimize_icon,
        command=app.iconify
    )
    minimize_button.place(relx=0.1, rely=0.025, anchor="nw")
    
    
    def show_details_frame():
        next_btn.place_forget()  
        main_frame.place_forget()
        details_frame.place(relx=0.5, rely=0.08, relwidth=0.9, relheight=0.80, anchor="n") 

    def back_to_main():
        details_frame.place_forget()
        main_frame.place(relx=0.5, rely=0.08, relwidth=0.9, anchor="n")
        next_btn.place(relx=0.5, rely=0.82, anchor="n", relwidth=0.9) 

        progress_bar.place_forget()
        status_label.lower()


    def finish_export():
        export_btn.configure(state="disabled")
        app.update()

        metadata = {
            "instructor": instructor_entry.get().strip(),
            "course": course_entry.get().strip(),
            "semester": semester_entry.get().strip(),
            "notes": notes_entry.get().strip()
        }

        selected_files = [path for var, path in file_vars if var.get()]
        author = author_entry.get().strip()
        out_folder = output_folder_entry.get().strip() or "."
        file_name = output_file_entry.get().strip()

        if not selected_files:
            show_confetti(success=False, message="NO FILES SELECTED")
            export_btn.configure(state="normal")
            app.update()
            return
    
        if not file_name.lower().endswith(".docx"):
            file_name += ".docx"

        output_path = os.path.join(out_folder, file_name)

        try:
            execution_outputs, visual_outputs = run_files_and_capture_output(selected_files)
            success = export_programs(author, selected_files, output_path, metadata, execution_outputs, visual_outputs)

            if success:
                show_confetti(success=True)
            else:
                show_confetti(success=False, message="CANNOT SAVE WORD FILE")
        except Exception as e:
            print(f"Export Failed: {e}")
            show_confetti(success=False, message="EXPORT FAILED")
        finally:
            export_btn.configure(state="normal")
            app.update()

    file_vars = []
    checkboxes = []

    # Title
    title_label = ctk.CTkLabel(app, text="AutoExport", font=("Segoe UI", 28, "bold"))
    title_label.place(relx=0.5, rely=0.015, anchor="n")

    # FRAME 1: MAIN INPUT
    main_frame = ctk.CTkFrame(app, corner_radius=12)
    main_frame.place(relx=0.5, rely=0.08, relwidth=0.9, anchor="n")

    def refresh_file_list(*args):
        global file_selection_state
        for widget in scroll_frame.winfo_children():
            widget.destroy()
        file_vars.clear()
        checkboxes.clear()

        folders = prog_folder_entry.get().strip().split(";")
        SUPPORTED_EXTS = (".py", ".java", ".cpp", ".c", ".js", ".go", ".rb", ".php", ".cs")

        files = []
        for folder in folders:
            folder = folder.strip()
            if os.path.isdir(folder):
                for root, _, files_in_dir in os.walk(folder):
                    for f in files_in_dir:
                        if f.endswith(SUPPORTED_EXTS):
                            path = os.path.join(root, f)
                            mod_time = os.path.getmtime(path)
                            size = os.path.getsize(path)
                            files.append((path, mod_time, size))

        sort_option = sort_selector.get() if 'sort_selector' in globals() else "Date Modified"

        if sort_option == "Date Modified":
            files.sort(key=lambda x: x[1], reverse=True)
        elif sort_option == "File Size":
            files.sort(key=lambda x: x[2], reverse=True)
        elif sort_option.startswith("Language: "):
            lang = sort_option.split(": ")[1].lower()
            ext_map = {
                "python": ".py", "java": ".java", "c++": ".cpp", "c": ".c",
                "javascript": ".js", "go": ".go", "ruby": ".rb", "php": ".php", "c#": ".cs"
            }
            ext = ext_map.get(lang)
            if ext:
                files = [f for f in files if f[0].lower().endswith(ext)]

        search_term = search_entry.get().lower()
        found_any = False

        for path, mod_time, _ in files:
            filename = os.path.basename(path)
            if search_term in filename.lower():
                found_any = True
                var = ctk.BooleanVar(value=file_selection_state.get(path, False))
                def make_callback(p=path, v=var):
                    def _cb(*_):
                        file_selection_state[p] = v.get()
                    return _cb

                var.trace_add("write", make_callback())

                cb = ctk.CTkCheckBox(
                    scroll_frame,
                    text=f"{filename} ({datetime.datetime.fromtimestamp(mod_time).strftime('%d %b %Y')})",
                    variable=var
                )
                cb.pack(anchor="w", padx=10, pady=2)
                file_vars.append((var, path))
                checkboxes.append(cb)

        if not found_any:
            empty_label = ctk.CTkLabel(scroll_frame, text="No files found", text_color="gray")
            empty_label.pack(pady=20)

        scroll_frame._parent_canvas.yview_moveto(0)

    def browse_folder(entry_widget):
        folder_path = filedialog.askdirectory()
        if folder_path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, folder_path)
            highlight_required(entry_widget, False)
            refresh_file_list()

    search_row = ctk.CTkFrame(main_frame, fg_color="transparent", border_width=0)
    toggle_btn = ctk.CTkButton(search_row, width=36, height=30, text="")
    attach_tooltip(toggle_btn,"Select/deselect all")

    def toggle_select_all():
        all_selected = all(var.get() for var, _ in file_vars)
        new_state = not all_selected
        for var, _ in file_vars:
            var.set(new_state)
        toggle_btn.configure(image=checkmark_icon if new_state else uncheck_icon)
        return new_state 

    def validate_and_show_details():
        author = author_entry.get().strip()
        prog_folders = prog_folder_entry.get().strip()
        output_folder = output_folder_entry.get().strip()
        file_name = output_file_entry.get().strip()

        def is_valid_author(name):
            return len(name) >= 4 and all(c.isalpha() or c.isspace() for c in name)

        def is_valid_output_file(name):
            return name.lower().endswith(".docx") and len(name) > 1

        def is_valid_folder_paths(paths):
            return any(os.path.isdir(path.strip()) for path in paths.split(";"))

        def is_valid_output_folder(folder):
            return folder == "" or os.path.isdir(folder)

        valid = True

        if not is_valid_author(author):
            highlight_required(author_entry, True)
            valid = False
        else:
            highlight_required(author_entry, False)

        if not is_valid_folder_paths(prog_folders):
            highlight_required(prog_folder_entry, True)
            valid = False
        else:
            highlight_required(prog_folder_entry, False)

        if not is_valid_output_file(file_name):
            highlight_required(output_file_entry, True)
            valid = False
        else:
            highlight_required(output_file_entry, False)

        if not is_valid_output_folder(output_folder):
            highlight_required(output_folder_entry, True)
            valid = False
        else:
            highlight_required(output_folder_entry, False)

        if not valid:
            return

        append_to_json("author", author), append_to_json("folder", prog_folders), append_to_json("outfolder", output_folder), append_to_json("output", file_name)
        show_details_frame()
    
    author_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Shresth Dwivedi", height=32)
    prog_folder_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Programs or path/to/code", height=32)
    prog_folder_entry.bind("<KeyRelease>", refresh_file_list)
    output_folder_entry = ctk.CTkEntry(main_frame, placeholder_text="Export folder (default: .)", height=32)
    output_file_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Assignment.docx", height=32)

    attach_tooltip(author_entry, "Enter your full name (at least 4 letters)")
    attach_tooltip(prog_folder_entry, "Enter location of programs")
    attach_tooltip(output_folder_entry, "Enter output-file's creation's location")
    attach_tooltip(output_file_entry, "Must end with .docx (e.g. Assignment.docx)")

    # Load saved values
    if author := config.get("author"):
        author_entry.insert(0, author)

    if folder := config.get("folder"):
        prog_folder_entry.insert(0, folder)

    if outfolder := config.get("outfolder"):
        output_folder_entry.insert(0, outfolder)

    if output := config.get("output"):
        output_file_entry.insert(0, output)

    # Enable live validation
    bind_validation(author_entry, "author")
    bind_validation(prog_folder_entry, "prog_folder")
    bind_validation(output_folder_entry, "output_folder")
    bind_validation(output_file_entry, "output_file")
    
    labels = ["Author Name", "Programs Folder(s)", "Export Folder", "Output File Name (.docx)"]
    entries = [author_entry, prog_folder_entry, output_folder_entry, output_file_entry]
    for i, (label, entry) in enumerate(zip(labels, entries)):
        ctk.CTkLabel(main_frame, text=label).grid(row=i, column=0, sticky="w", padx=20, pady=(8 if i == 0 else 8))
        entry.grid(row=i, column=1, padx=20, pady=(8 if i == 0 else 8), sticky="ew")
        if "Folder" in label:
            ctk.CTkButton(main_frame, width=36, height=30, text="", image=folder_icon, command=lambda e=entry: browse_folder(e)).grid(row=i, column=2, padx=(0,20), sticky="ew")

    author_entry.grid(row=0, column=1, columnspan=2, padx=20, pady=(15 if i == 0 else 8), sticky="ew")
    output_file_entry.grid(row=3, column=1, columnspan=2, padx=20, pady=(15 if i == 0 else 8), sticky="ew")

    main_frame.grid_columnconfigure(1, weight=1)
    
    search_row.grid(row=4, column=0, columnspan=3, padx=10, pady=(10, 5), sticky="ew")
    search_row.grid_columnconfigure(0, weight=1) 

    search_entry = ctk.CTkEntry(search_row, placeholder_text="Search filename...", width=499)
    search_entry.grid(row=4, column=0, columnspan=2, padx=(10), pady=(10, 5), sticky="w")
    search_entry.bind("<KeyRelease>", refresh_file_list)

    initial_all_selected = all(var.get() for var, _ in file_vars)
    toggle_btn.configure(image=uncheck_icon if initial_all_selected else checkmark_icon, compound="left", command=toggle_select_all)
    toggle_btn.grid(row=4, column=2, padx=(0,10), pady=(10, 5), sticky="ew")
    
    sort_options = [
        "Date Modified", "File Size",
        "Language: Python", "Language: Java", "Language: C++", "Language: C", "Language: JavaScript", 
        "Language: Go", "Language: Ruby", "Language: PHP", "Language: C#"
    ]
    sort_selector = ctk.CTkOptionMenu(
        master=search_row,
        values=sort_options,
        width=163,
        command=refresh_file_list
    )
    sort_selector.set("Date Modified")
    sort_selector.grid(row=4, column=1, columnspan=1, padx=20, pady=(10, 5), sticky="e")

    container_frame = ctk.CTkFrame(main_frame)
    container_frame.grid(row=5, column=0, columnspan=3, padx=20, pady=(5, 15), sticky="nsew")
    scroll_frame = ctk.CTkScrollableFrame(container_frame, height=200, border_width=0, fg_color="transparent")
    scroll_frame.pack(fill="both", expand=True, padx=(5, 5), pady=(5, 5)) 

    next_btn = ctk.CTkButton(app, text="Next", fg_color="green", text_color="white", height=45, command=validate_and_show_details)
    next_btn.place(relx=0.5, rely=0.82, anchor="n", relwidth=0.9)

    progress_bar = ctk.CTkProgressBar(app, mode="determinate")
    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.set(0)
    progress_bar.configure(determinate_speed=1)
    progress_bar.place_forget()

    # FRAME 2: DETAILS FRAME
    details_frame = ctk.CTkFrame(app, corner_radius=12)

    # gif_label = ctk.CTkLabel(details_frame, text="", fg_color="transparent")
    # gif_label.place(relx=0.99, rely=1, anchor="se")

    status_label = ctk.CTkLabel(details_frame, text="", font=("Segoe UI", 11))
    status_label.place(relx=0.03, rely=0.996, anchor="sw") 
    #status_label.lower() 

    instructor_entry = ctk.CTkEntry(details_frame, placeholder_text="Instructor's Name (optional)", height=32)
    course_entry = ctk.CTkEntry(details_frame, placeholder_text="Course / Subject Name (optional)", height=32)
    semester_entry = ctk.CTkEntry(details_frame, placeholder_text="Semester or Session (optional)", height=32)
    notes_entry = ctk.CTkEntry(details_frame, placeholder_text="Additional Notes (optional)", height=32)

    meta_fields = [("Instructor", instructor_entry),
                   ("Course / Subject", course_entry),
                   ("Semester", semester_entry),
                   ("Notes", notes_entry)]

    for i, (label, entry) in enumerate(meta_fields):
        ctk.CTkLabel(details_frame, text=label).grid(row=i, column=0, sticky="w", padx=20, pady=(8 if i == 0 else 8))
        entry.grid(row=i, column=1, padx=20, pady=(8 if i == 0 else 8), sticky="ew")

    details_frame.grid_columnconfigure(1, weight=1)

    export_btn = ctk.CTkButton(details_frame, text="Export", fg_color="green", text_color="white", height=45, command=lambda: threading.Thread(target=finish_export, daemon=True).start())
    export_btn.grid(row=6, column=0, columnspan=2, padx=20, pady=(10), sticky="ew")

    back_btn = ctk.CTkButton(details_frame, text="Back", image=back_icon, command=back_to_main)
    back_btn.grid(row=7, column=0, columnspan=2, padx=20, pady=(0), sticky="ew")

    def toggle_theme_mode():
        current_mode = ctk.get_appearance_mode()
        new_mode = "light" if current_mode == "Dark" else "dark"
        
        ctk.set_appearance_mode(new_mode)
        
        new_icon = dark_theme_icon if new_mode == "light" else light_theme_icon
        theme_mode_button.configure(image=new_icon)

        github_btn.configure(image=(github_light_icon if new_mode == "dark" else github_dark_icon))
        linkedin_btn.configure(image=(linkedin_light_icon if new_mode == "dark" else linkedin_dark_icon))
        x_btn.configure(image=(X_light_icon if new_mode == "dark" else X_dark_icon))
        bluesky_btn.configure(image=(bluesky_light_icon if new_mode == "dark" else bluesky_dark_icon))
        youtube_btn.configure(image=(youtube_light_icon if new_mode == "dark" else youtube_dark_icon))

        append_to_json("mode", new_mode)
        app.update_idletasks()

    theme_mode_button = ctk.CTkButton(app, width=36, height=30, text="",image=(light_theme_icon if load_mode() == "dark" else dark_theme_icon), command=toggle_theme_mode)
    theme_mode_button.place(relx=0.95, rely=0.025, anchor="ne")
    attach_tooltip(theme_mode_button,"Toggle light/dark mode")

    def toggle_theme():
        theme_selection_popup()

    def get_current_theme_path():
        try:
            with open(".autoexport/cache.json", "r") as f:
                cache = json.load(f)
                return cache.get("theme")
        except Exception as e:
            #print(f"Error reading theme path: {e}")
            return None

    theme_button = ctk.CTkButton(app, width=36, height=30, text="", image=theme_icon, command=toggle_theme)
    theme_button.place(relx=0.9, rely=0.025, anchor="ne")
    attach_tooltip(theme_button, "Change app's theme")

    donation_frame = ctk.CTkFrame(details_frame, corner_radius=10)
    donation_frame.grid(row=8, column=0, columnspan=2, pady=(30, 10), padx=20, sticky="ew")

    button_theme = ctk.ThemeManager.theme.get("CTkButton", {})
    default_fg_color = button_theme.get("fg_color", ["#3B8ED0", "#1F6AA5"])
    default_hover_color = button_theme.get("hover_color", ["#36719F", "#144870"])
   
    theme_path = get_current_theme_path() 

    def get_theme_font_from_json(theme_path: str):
        try:
            with open(theme_path, 'r', encoding='utf-8') as f:
                theme_data = json.load(f)

            if platform.system() == "Windows":
                font_info = theme_data.get("CTkFont", {}).get("Windows", {})

            family = font_info.get("family", "Segoe UI")
            size = int(font_info.get("size", 13))
            return family, size, "italic"

        except Exception as e:
            return "Segoe UI", 13, "italic"

    support_label_font = get_theme_font_from_json(theme_path)
    
    support_label = ctk.CTkLabel(
        donation_frame,
        text="Support via UPI or follow me here:",
        font=support_label_font,
        anchor="center"
    )
    support_label.pack(padx=10, pady=(10, 5))

    upi_id_label = ctk.CTkLabel(
        donation_frame,
        text="shresthdwivedi03@axl",
        font=("JetBrains Mono", 11),
        #text_color="lightgreen"
    )
    upi_id_label.pack(pady=(0, 5))
    
    def copy_upi_to_clipboard():
        app.clipboard_clear()
        app.clipboard_append("shresthdwivedi03@axl")
        app.update()

        copy_btn.configure(text=" Copied!", image=checkmark_icon, fg_color="green", hover_color="green")
        app.after(2000, lambda: copy_btn.configure(
            text=" Copy UPI",
            image=clipboard_icon,
            fg_color=default_fg_color,
            hover_color=default_hover_color
        ))

    copy_btn = ctk.CTkButton(
        donation_frame,
        text=" Copy UPI",
        image=clipboard_icon,
        fg_color=default_fg_color,
        hover_color=default_hover_color,
        command=copy_upi_to_clipboard
    )
    copy_btn.pack(pady=(0, 10))

    attach_tooltip(copy_btn, "Click to copy UPI ID")

    social_frame = ctk.CTkFrame(donation_frame, fg_color="transparent", border_width=0)
    social_frame.pack(pady=(5, 10))

    current_mode = ctk.get_appearance_mode()
    
    # GitHub icon button
    github_btn = ctk.CTkLabel(social_frame, text="", image=(github_light_icon if current_mode == "Dark" else github_dark_icon), cursor="hand2")
    github_btn.pack(side="left", padx=20, pady=(1,1))
    github_btn.bind("<Button-1>", lambda e: webbrowser.open("https://github.com/Shresth-Dwivedi"))

    # LinkedIn icon button
    linkedin_btn = ctk.CTkLabel(social_frame, text="", image=(linkedin_light_icon if current_mode == "Dark" else linkedin_dark_icon), cursor="hand2")
    linkedin_btn.pack(side="left", padx=20, pady=(1,1))
    linkedin_btn.bind("<Button-1>", lambda e: webbrowser.open("https://www.linkedin.com/in/shresth-dwivedi/"))

    # Twitter icon button
    x_btn = ctk.CTkLabel(social_frame, text="", image=(X_light_icon if current_mode == "Dark" else X_dark_icon), cursor="hand2")
    x_btn.pack(side="left", padx=20, pady=(1,1))
    x_btn.bind("<Button-1>", lambda e: webbrowser.open("https://x.com/theDavyDee"))

    # Bluesky icon button
    bluesky_btn = ctk.CTkLabel(social_frame, text="", image=(bluesky_light_icon if current_mode == "Dark" else bluesky_dark_icon), cursor="hand2")
    bluesky_btn.pack(side="left", padx=20, pady=(1,1))
    bluesky_btn.bind("<Button-1>", lambda e: webbrowser.open("https://bsky.app/profile/shresthdwivedi.bsky.social"))
    
    # Youtube icon button
    youtube_btn = ctk.CTkLabel(social_frame, text="", image=(youtube_light_icon if current_mode == "Dark" else youtube_dark_icon), cursor="hand2")
    youtube_btn.pack(side="left", padx=20, pady=(1,1))
    youtube_btn.bind("<Button-1>", lambda e: webbrowser.open("http://www.youtube.com/@shresth.dwivedi"))

    credit_frame = ctk.CTkFrame(app, fg_color="transparent", border_width=0)
    credit_frame.place(relx=0.5, rely=0.9, anchor="n")
    
    # Copyright label
    copyright_label = ctk.CTkLabel(
        credit_frame,
        text="© 2025 Shresth Dwivedi. All rights reserved.",
        font=("Segoe UI", 11)
    )
    copyright_label.pack()

    # Version label
    version_label = ctk.CTkLabel(
        app,
        text="2.3 64-bit",
        font=("Segoe UI", 11)
    )
    version_label.place(relx=1.0, rely=1.0, anchor="se", x=-42, y=-18)

    refresh_file_list()

    def show_update_popup():
        update_msg = check_for_update("2.3")
        if not update_msg:
            return

        popup = ctk.CTkButton(
            app,
            text=update_msg,
            image=bell_icon,
            text_color="white",
            fg_color="orange",
            hover_color="orange",
            compound="left",
            font=("Segoe UI", 11),
            cursor="hand2"
        )

        popup.place(relx=0.5, rely=-0.1, anchor="n") 

        def open_release(event=None):
            webbrowser.open_new_tab("https://github.com/Shresth-Dwivedi/AutoExport/releases/latest")

        popup.bind("<Button-1>", open_release)

        def animate(step=0, reverse=False):
            if reverse:
                new_y = 0.07 - (step * 0.07)
            else:
                new_y = -0.1 + (step * 0.07)

            popup.place_configure(rely=new_y)

            if (not reverse and new_y < 0.07) or (reverse and new_y > -0.1):
                app.after(30, lambda: animate(step + 0.1, reverse))
            elif not reverse:
                app.after(10000, lambda: animate(0, reverse=True))
            else:
                app.after(60000, lambda: try_show_popup())

        def try_show_popup():
            if is_internet_available():
                animate()

        app.after(5000, try_show_popup)

    show_update_popup()

    def open_terms():
        webbrowser.open("https://github.com/Shresth-Dwivedi/AutoExport/blob/main/LICENSE") 

    terms_label = ctk.CTkLabel(
        credit_frame,
        text="Terms & Conditions",
        cursor="hand2",
        font=ctk.CTkFont("Segoe UI", 11, underline=True)
    )
    terms_label.pack()
    terms_label.bind("<Button-1>", lambda e: open_terms())

    # Character animation state
    character_state = {"frames": [], "index": 0, "label": None, "animation": "idle"}

    def load_animation(path, scale=3, flip=True):
        frames = []
        try:
            with Image.open(path) as im:
                for i in range(im.n_frames):
                    im.seek(i)
                    frame = im.copy()
                    if flip:
                        frame = frame.transpose(Image.FLIP_LEFT_RIGHT)
                    new_size = (frame.width * scale, frame.height * scale)
                    frames.append(
                        ctk.CTkImage(
                            light_image=frame.resize(new_size, Image.NEAREST),
                            dark_image=frame.resize(new_size, Image.NEAREST),
                            size=new_size
                        )
                    )
        except Exception as e:
            print(f"[!] Could not load {path}: {e}")
        return frames

    # Load both animations
    idle_frames = load_animation(resource_path("assets/buddy/idle-2.gif"))
    wave_frames = load_animation(resource_path("assets/buddy/idle.gif"))

    # Set initial animation
    character_state["frames"] = idle_frames

    character_state["label"] = ctk.CTkLabel(details_frame, text="", image=idle_frames[0])
    character_state["label"].place(relx=0.99, rely=1, anchor="se")

    woof_label = ctk.CTkLabel(details_frame, text="", image=woof_img)
    woof_label.place(relx=0.88, rely=0.95, anchor="se")
    woof_label.lower()

    def animate_character():
        frames = character_state["frames"]
        label = character_state["label"]

        if not frames:
            return

        character_state["index"] = (character_state["index"] + 1) % len(frames)
        label.configure(image=frames[character_state["index"]])
        app.after(100, animate_character)

    def switch_animation(event=None):
        if character_state["animation"] != "wave":
            woof_label.lift()
            character_state["frames"] = wave_frames
            character_state["animation"] = "wave"
            character_state["index"] = 0

            app.after(5000, revert_to_idle)

    def revert_to_idle():
        woof_label.lower()
        character_state["frames"] = idle_frames
        character_state["animation"] = "idle"
        character_state["index"] = 0

    character_state["label"].bind("<Button-1>", switch_animation)

    animate_character()

    app.mainloop()

if __name__ == "__main__":
    modern_gui()