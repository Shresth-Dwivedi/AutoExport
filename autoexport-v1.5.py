import os
import io
import sys
import json
import shutil
import datetime
import subprocess
import traceback
import threading
import urllib.request
import socket
import webbrowser
import customtkinter as ctk
import threading, queue, time
from subprocess import Popen, PIPE
from PIL import Image
from tkinter import filedialog
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

CONFIG_FILE = "settings.json"
DEFAULT_FILE_NAME = "Assignment.docx"

def is_internet_available():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=2)
        return True
    except OSError:
        return False


def check_for_update(local_version="2.0.0"):
    if not is_internet_available():
        return None

    try:
        url = "https://raw.githubusercontent.com/Shresth-Dwivedi/AutoExport/main/version.txt"
        with urllib.request.urlopen(url) as response:
            latest_version = response.read().decode("utf-8").strip()

        if latest_version != local_version:
            return f"Update available: v{latest_version}"
        return None
    except Exception:
        return None

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller bundled .exe """
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {}

def save_config(data):
    with open(CONFIG_FILE, "w") as f:
        json.dump(data, f)

def get_sorted_file_paths(paths, sort_by):
        if sort_by == "Date Created":
            return sorted(paths, key=lambda p: os.path.getctime(p))
        elif sort_by == "File Size":
            return sorted(paths, key=lambda p: os.path.getsize(p))
        elif sort_by == "Language":
            return sorted(paths, key=lambda p: os.path.splitext(p)[1])
        return paths

def export_programs(author, selected_files, output_path, metadata=None, execution_outputs=None, visual_outputs=None):
    if not selected_files:
        return "No files selected for export."

    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.lift()
    progress_bar.set(0)
    progress_bar.update()
    status_label.configure(text="Starting export...")
    status_label.lift()
    try:
        if app.winfo_exists():
            app.update_idletasks()
    except:
        pass

    doc = Document()

    if 'CodeBlock' not in doc.styles:
        style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.right_indent = Inches(0.3)

    doc.add_heading("Programming Assignment", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Author: {author}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Exported on: {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if metadata:
        doc.add_paragraph(f"Instructor: {metadata.get('instructor')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Course: {metadata.get('course')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Semester: {metadata.get('semester')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if metadata.get("notes"):
            doc.add_paragraph(f"Notes: {metadata.get('notes')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()
    doc.add_paragraph("Table of Contents", style="Heading 1")
    doc.add_paragraph("→ Update fields in Word via References → Table of Contents → Update Field.")
    doc.add_page_break()

    # Apply sorting before export
    selected_sort = sort_selector.get()
    selected_files = get_sorted_file_paths(selected_files, selected_sort)

    total_files = len(selected_files)

    for i, path in enumerate(selected_files, 1):
        try:
            filename = os.path.basename(path)
            status_label.configure(text=f"Exporting...")
            app.update_idletasks()

            with open(path, "r", encoding="utf-8") as f:
                content = f.read()

            modified_time = datetime.datetime.fromtimestamp(os.path.getmtime(path)).strftime("%d %b %Y, %H:%M")

            doc.add_heading(filename, level=2)
            doc.add_paragraph(f"Last Modified: {modified_time}", style="Normal")

            for line in content.splitlines():
                doc.add_paragraph(line, style='CodeBlock')

            #status_label.configure(text=f"Writing output for: {filename}")
            #app.update_idletasks()

            doc.add_paragraph("Output:", style="Heading 2")
            output = (
                execution_outputs.get(path)
                or execution_outputs.get(filename)
                or "Output not available."
            )

            for line in output.splitlines():
                doc.add_paragraph(line, style='CodeBlock')

            if visual_outputs:
                images = visual_outputs.get(path, []) or visual_outputs.get(filename, [])
                if images:
                    #status_label.configure(text=f"Adding visuals for: {filename}")
                    #app.update_idletasks()

                    doc.add_paragraph("Visual Output(s):", style="Heading 2")
                    for image_path in images:
                        if os.path.exists(image_path):
                            try:
                                doc.add_picture(image_path, width=Inches(5.5))
                                doc.add_paragraph(f"[{os.path.basename(image_path)}]", style='Normal')
                            except Exception as e:
                                doc.add_paragraph(f"[Failed to load image: {image_path} — {e}]", style='Normal')
                        else:
                            doc.add_paragraph(f"[Missing image file: {image_path}]", style='Normal')

        except Exception as e:
            print(f"Skipped file {path}: {e}")
            doc.add_paragraph(f"[Failed to process file: {path}]", style="Normal")

        doc.add_page_break()

        progress = i / total_files
        progress_bar.set(progress)
        try:
            if app.winfo_exists():
                app.update_idletasks()
        except:
            pass

    status_label.configure(text="Saving document...")
    app.update_idletasks()

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Failed to save Word file: {e}")
         # Hide progress bar and status label
        if app.winfo_exists():
            status_label.configure(text="")
            progress_bar.lower()
            return False  # failed

    try:
        status_label.configure(text="Opening exported file...")
        app.update_idletasks()
        subprocess.Popen(["start", "", output_path], shell=True)
    except Exception as e:
        print(f"Could not open Word file automatically: {e}")

    if app.winfo_exists():
        status_label.configure(text="")
        app.after(1000, lambda: progress_bar.lower())

    return True  # success

def terminal_input_session(program_name, input_prompts):
    inputs = []

    dialog = ctk.CTkToplevel()
    dialog.title(f"Terminal Input – {program_name}")
    dialog.geometry("600x400")
    dialog.resizable(False, False)

    # Center dialog
    dialog.update_idletasks()
    w, h = 600, 400
    x = (dialog.winfo_screenwidth() // 2) - (w // 2)
    y = (dialog.winfo_screenheight() // 2) - (h // 2)
    dialog.geometry(f"{w}x{h}+{x}+{y}")
    dialog.grab_set()

    output_box = ctk.CTkTextbox(dialog, height=280, wrap="word", font=("Consolas", 12))
    output_box.pack(fill="both", expand=True, padx=10, pady=(10, 5))
    output_box.configure(state="disabled")

    entry_var = ctk.StringVar()
    entry = ctk.CTkEntry(dialog, textvariable=entry_var, font=("Consolas", 12))
    entry.pack(fill="x", padx=10, pady=(0, 10))
    entry.focus_set()

    current_index = [0]

    def print_to_terminal(text):
        output_box.configure(state="normal")
        output_box.insert("end", text + "\n")
        output_box.configure(state="disabled")
        output_box.see("end")

    def on_submit(*_):
        if current_index[0] < len(input_prompts):
            answer = entry_var.get()
            inputs.append(answer)

            print_to_terminal(f"{input_prompts[current_index[0]]}{answer}")
            entry_var.set("")
            current_index[0] += 1

            if current_index[0] < len(input_prompts):
                print_to_terminal(input_prompts[current_index[0]])
            else:
                dialog.destroy()

    entry.bind("<Return>", on_submit)

    print_to_terminal(f"Running: {program_name}")
    print_to_terminal(input_prompts[0])

    dialog.wait_window()
    return inputs

def detect_language_and_command(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    filedir = os.path.dirname(file_path)
    basename = os.path.splitext(filename)[0]

    if ext == ".py" and shutil.which("python"):
        return "Python", ["python", file_path], None

    elif ext == ".cpp" and shutil.which("g++"):
        output = os.path.join(filedir, f"{basename}.exe" if os.name == "nt" else basename)
        return "C++", ["g++", file_path, "-o", output], output

    elif ext == ".c" and shutil.which("gcc"):
        output = os.path.join(filedir, f"{basename}.exe" if os.name == "nt" else basename)
        return "C", ["gcc", file_path, "-o", output], output

    elif ext == ".java" and shutil.which("javac") and shutil.which("java"):
        return "Java", ["javac", file_path], basename

    elif ext == ".js" and shutil.which("node"):
        return "JavaScript", ["node", file_path], None

    elif ext == ".go" and shutil.which("go"):
        return "Go", ["go", "run", file_path], None

    elif ext == ".rb" and shutil.which("ruby"):
        return "Ruby", ["ruby", file_path], None

    elif ext == ".php" and shutil.which("php"):
        return "PHP", ["php", file_path], None

    elif ext == ".cs" and shutil.which("csc"):
        exe_name = os.path.join(filedir, f"{basename}.exe")
        return "C#", ["csc", "/out:" + exe_name, file_path], exe_name

    return None, [], None

def create_terminal_window(title="Custom Terminal"):

    dialog = ctk.CTkToplevel(master=app)
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(__file__)

    app_icon = os.path.join(base_path, "icons", "app.ico")
    dialog.iconbitmap(app_icon)
    dialog.protocol("WM_DELETE_WINDOW", lambda: None)
    dialog.title(title)
    dialog.geometry("650x400")
    dialog.resizable(False, False)

    dialog.update_idletasks()
    w, h = 650, 400
    x = (dialog.winfo_screenwidth() // 2) - (w // 2)
    y = (dialog.winfo_screenheight() // 2) - (h // 2)
    dialog.geometry(f"{w}x{h}+{x}+{y}")
    terminal_closed = [False] 
    dialog.grab_set()

    output_box = ctk.CTkTextbox(dialog, font=("Consolas", 12), wrap="word")
    output_box.pack(fill="both", expand=True, padx=10, pady=(10, 5))
    output_box.configure(state="disabled")

    input_var = ctk.StringVar()
    input_entry = ctk.CTkEntry(dialog, textvariable=input_var, font=("Consolas", 12))
    input_entry.pack(fill="x", padx=10, pady=(0, 10))

    prompt_event = threading.Event()
    input_queue = queue.Queue()
    collected_output = []  
    
    def write(text):
        output_box.configure(state="normal")
        output_box.insert("end", text)
        output_box.configure(state="disabled")
        output_box.see("end")
        collected_output.append(text)   

    def wait_for_prompt(tag="input"):
        def enable_input():
            prompt_event.set()
            input_entry.focus_set()

        dialog.after(100, enable_input)
        
    def get_next_input():
        prompt_event.wait(100)  
        value = input_queue.get()  
        prompt_event.clear()
        return value

    def on_submit(event=None):
        # if not prompt_event.is_set():
        #     return
        value = input_var.get()
        input_var.set("")
        input_queue.put(value)
        write(value + "\n")  
        
    input_entry.bind("<Return>", on_submit)

    def destroy():
        terminal_closed[0] = True
        try:
            dialog.destroy()
        except:
            pass

    def is_closed():
        return terminal_closed[0]

    def exists():
        return bool(dialog.winfo_exists())

    return type("Terminal", (), {
        "write": staticmethod(write),
        "wait_for_prompt": staticmethod(wait_for_prompt),
        "get_next_input": staticmethod(get_next_input),
        "destroy": staticmethod(destroy),
        "is_closed": staticmethod(is_closed),
        "exists": staticmethod(exists),
        "collected_output": collected_output,
    })()


def execute_with_input_detection(cmd_list, program_name, force_terminal=False):
    
    q = queue.Queue()
    collected_output = []
    buffer = ""
    terminal = None
    show_terminal = False

    if force_terminal:
        terminal = create_terminal_window(program_name)
        show_terminal = True

    process = Popen(cmd_list, stdin=PIPE, stdout=PIPE, stderr=PIPE, text=True, bufsize=0)

    def read_stream(stream, tag):
        for chunk in iter(lambda: stream.read(1), ''):
            q.put((tag, chunk))
        q.put((tag, None))

    threading.Thread(target=read_stream, args=(process.stdout, "stdout"), daemon=True).start()
    threading.Thread(target=read_stream, args=(process.stderr, "stderr"), daemon=True).start()

    last_output_time = time.time()
    silence_threshold = 1.5
    show_terminal = force_terminal
    input_waiting = False

    while True:
        try:
            if terminal and terminal.is_closed():
                try:
                    process.terminate()
                except:
                    pass
                return "[⚠] Terminal was closed by user. Skipped program.\n"

            while not q.empty():
                tag, chunk = q.get_nowait()
                if chunk is None:
                    continue

                collected_output.append(chunk)
                buffer += chunk

                if show_terminal:
                    terminal.write(chunk if tag == "stdout" else f"[stderr] {chunk}")
                elif not force_terminal and (time.time() - last_output_time > silence_threshold):
                    terminal = create_terminal_window(program_name)
                    terminal.write("".join(collected_output))
                    show_terminal = True
                    buffer = ""

                last_output_time = time.time()
                input_waiting = False

            if show_terminal and not input_waiting and (time.time() - last_output_time > silence_threshold):
                input_waiting = True
                terminal.wait_for_prompt("input")
                user_input = terminal.get_next_input()
                if user_input is None:
                    break
                process.stdin.write(user_input + "\n")
                process.stdin.flush()

                collected_output.append(user_input + "\n")

                last_output_time = time.time()
                input_waiting = False

            if process.poll() is not None:
                break

        except Exception:
            break

    process.wait()

    if show_terminal and terminal:
        terminal.write("\n[✔] Program finished. Press Enter to close...\n")
        terminal.wait_for_prompt("exit")
        _ = terminal.get_next_input()
        terminal.destroy()

    return "".join(collected_output)

def program_requires_input(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read().lower()
            ext = os.path.splitext(file_path)[1].lower()

            input_indicators = []

            if ext == ".py":
                input_indicators = [
                    "input(", "sys.stdin.read", "sys.stdin.readline", "raw_input(",
                ]

            elif ext == ".java":
                input_indicators = [
                    "scanner", ".next", "system.in", "bufferedreader", ".readline",
                    "datainputstream", ".read(", ".readline("
                ]

            elif ext in [".c", ".cpp"]:
                input_indicators = [
                    "scanf", "getchar", "fgets", "gets(", "cin", "std::cin", "stdin",
                    "read(", "fgetc(", "getch("
                ]

            elif ext == ".cs":
                input_indicators = [
                    "console.readline", "streamreader", "textreader", ".readline", ".read("
                ]

            elif ext == ".js":
                input_indicators = [
                    "prompt(", "readline(", "process.stdin", "fs.createinterfacestream", "require('readline')"
                ]

            elif ext == ".go":
                input_indicators = [
                    "fmt.scan", "fmt.scanln", "bufio", "os.stdin", ".readstring"
                ]

            elif ext == ".rb":
                input_indicators = [
                    "gets", "stdin.gets", "readline"
                ]

            elif ext == ".php":
                input_indicators = [
                    "fgets(stdin", "readline(", "stream_socket_client(\"php://stdin", "$stdin"
                ]

            return any(keyword in code for keyword in input_indicators)

    except Exception:
        return False

            
def run_files_and_capture_output(file_paths):
    output_log = {}
    visual_outputs = {}

    VISUAL_TEMP_DIR = "generated_visuals"
    os.makedirs(VISUAL_TEMP_DIR, exist_ok=True)

    total_files = len(file_paths)

    # Show and reset progress bar
    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.lift()
    progress_bar.set(0.0)
    app.update_idletasks()

    for i, path in enumerate(file_paths):
        buffer = io.StringIO()
        filename = os.path.basename(path)
        visuals_for_file = []

        status_label.configure(text=f"Processing {filename}")
        app.update_idletasks()

        lang, compile_cmd, runtime_target = detect_language_and_command(path)

        if not lang:
            output_log[path] = "⚠️ Unsupported file type or required toolchain not found."
            visual_outputs[path] = []
            continue

        try:
            try:
                import matplotlib
                matplotlib.use("Agg")
                import matplotlib.pyplot as plt

                original_show = plt.show
                original_savefig = plt.savefig

                def patched_show(*args, **kwargs):
                    nonlocal visuals_for_file
                    fname = f"{filename}_auto_figure_{len(visuals_for_file)+1}.png"
                    save_path = os.path.join(VISUAL_TEMP_DIR, fname)
                    visuals_for_file.append(save_path)
                    original_savefig(save_path, *args, **kwargs)
                    plt.clf()
                    plt.cla()

                def patched_savefig(fname=None, *args, **kwargs):
                    if not fname:
                        fname = f"{filename}_figure_{len(visuals_for_file)+1}.png"
                    else:
                        fname = os.path.basename(fname)
                    save_path = os.path.join(VISUAL_TEMP_DIR, fname)
                    visuals_for_file.append(save_path)
                    original_savefig(save_path, *args, **kwargs)

                plt.show = patched_show
                plt.savefig = patched_savefig
            except ImportError:
                plt = None

            requires_input = program_requires_input(path)

            if lang == "Python":
                status_label.configure(text=f"Running {filename}")
                app.update_idletasks()
                if requires_input:
                    output = execute_with_input_detection(["python", path], filename, force_terminal=True)
                    buffer.write(output)
                else:
                    global_namespace = {"__name__": "__main__"}
                    with open(path, "r", encoding="utf-8") as f:
                        code = f.read()
                    with redirect_stdout(buffer):
                        exec(compile(code, path, 'exec'), global_namespace)

            elif lang == "Java":
                status_label.configure(text=f"⚙ Compiling {filename}")
                app.update_idletasks()
                compile_proc = subprocess.run(compile_cmd, capture_output=True, text=True)
                if compile_proc.stderr:
                    buffer.write("Compile error:\n" + compile_proc.stderr)
                else:
                    status_label.configure(text=f"Running {filename}")
                    app.update_idletasks()
                    class_dir = os.path.dirname(path)
                    main_class = runtime_target
                    output = execute_with_input_detection(["java", "-cp", class_dir, main_class], filename, force_terminal=requires_input)
                    buffer.write(output)

            elif lang in ["C", "C++", "C#"]:
                status_label.configure(text=f"⚙ Compiling {filename}")
                app.update_idletasks()
                subprocess.run(compile_cmd, capture_output=True)
                if runtime_target and os.path.exists(runtime_target):
                    status_label.configure(text=f"Running {filename}")
                    app.update_idletasks()
                    output = execute_with_input_detection([runtime_target], filename, force_terminal=requires_input)
                    buffer.write(output)
                else:
                    buffer.write("⚠️ Compiled binary not found.\n")

            elif lang in ["JavaScript", "Go", "Ruby", "PHP"]:
                status_label.configure(text=f"Running {filename}")
                app.update_idletasks()
                output = execute_with_input_detection(compile_cmd, filename, force_terminal=requires_input)
                buffer.write(output)

            else:
                buffer.write("⚠️ Unknown language execution path.\n")

        except Exception:
            buffer.write("Error during execution:\n")
            buffer.write(traceback.format_exc())

        if plt:
            plt.show = original_show
            plt.savefig = original_savefig

        output_log[path] = buffer.getvalue().strip()
        visual_outputs[path] = visuals_for_file

        # Update progress
        progress = (i + 1) / total_files
        progress_bar.set(progress)
        app.update_idletasks()

    #app.after(1000, lambda: progress_bar.lower())

    return output_log, visual_outputs


def attach_tooltip(widget, text, delay=1500):
    tooltip = None
    after_id = None

    def show_tooltip():
        nonlocal tooltip
        if not tooltip:
            tooltip = ctk.CTkLabel(widget.master, text=text, font=("Segoe UI", 10),
                                   text_color="white", bg_color="#333", corner_radius=4)
            x = widget.winfo_x()
            y = widget.winfo_y() - 25
            tooltip.place(x=x, y=y)

    def on_enter(event):
        nonlocal after_id
        after_id = widget.after(delay, show_tooltip)

    def on_leave(event):
        nonlocal tooltip, after_id
        if after_id:
            widget.after_cancel(after_id)
            after_id = None
        if tooltip:
            tooltip.destroy()
            tooltip = None

    widget.bind("<Enter>", on_enter)
    widget.bind("<Leave>", on_leave)


def bind_validation(entry, field_name):
    def validate(_event):
        value = entry.get().strip()

        def is_valid_author(name):
            return len(name) >= 4 and all(c.isalpha() or c.isspace() for c in name)

        def is_valid_output_file(name):
            return name.lower().endswith(".docx") and len(name) > 5

        def is_valid_folder_paths(paths):
            return any(os.path.isdir(path.strip()) for path in paths.split(";"))

        def is_valid_output_folder(folder):
            return folder == "" or os.path.isdir(folder)

        
        if field_name == "author":
            is_valid = is_valid_author(value)
        elif field_name == "prog_folder":
            is_valid = is_valid_folder_paths(value)
        elif field_name == "output_folder":
            is_valid = is_valid_output_folder(value)
        elif field_name == "output_file":
            is_valid = is_valid_output_file(value)
        else:
            is_valid = bool(value)  # Fallback for unexpected case

        highlight_required(entry, error=not is_valid)
    entry.bind("<KeyRelease>", validate)


def highlight_required(entry, error: bool = True):
    border_color = "red" if error else "green"
    entry.configure(border_color=border_color, border_width=2)

def show_confetti(success=True, message="Exported Successfully"):
    color = "green" if success else "red"
    #emoji = "🎉" if success else "❌"
    confetti = ctk.CTkLabel(app, text=f"{message}", font=("Segoe UI", 20), text_color=color)
    confetti.place(relx=0.5, rely=0.85, anchor="center")
    app.after(3000, confetti.destroy)

def modern_gui():
    global app, progress_bar, status_label, sort_selector
    config = load_config()
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")

    app = ctk.CTk()
    app.title("AutoExport")

    #icon for the app
    icon_path = resource_path(os.path.join("icons", "app.ico"))
    app.iconbitmap(default=icon_path)

    # Desired window size
    window_width = 850
    window_height = 680

    # Screen dimensions
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()

    vertical_offset = 40

    x = (screen_width // 2) - (window_width // 2)
    y = max(0, (screen_height // 2) - (window_height // 2) - vertical_offset)

    # Apply final geometry
    app.geometry(f"{window_width}x{window_height}+{x}+{y}")
    app.resizable(False, False)

    # Icons for labels and buttons
    bell_icon = ctk.CTkImage(Image.open(resource_path("icons/bell.png")), size=(16, 16))
    clipboard_icon = ctk.CTkImage(Image.open(resource_path("icons/clip.png")), size=(16, 16))
    checkmark_icon = ctk.CTkImage(Image.open(resource_path("icons/check.png")), size=(16, 16))
    back_icon = ctk.CTkImage(Image.open(resource_path("icons/back.png")), size=(16, 16))
    folder_icon = ctk.CTkImage(Image.open(resource_path("icons/open-folder.png")), size=(16, 16))

    # Social icons
    github_icon = ctk.CTkImage(Image.open(resource_path("icons/gitHub.png")), size=(24, 24))
    linkedin_icon = ctk.CTkImage(Image.open(resource_path("icons/linkedin.png")), size=(24, 24))
    X_icon = ctk.CTkImage(Image.open(resource_path("icons/X.png")), size=(24, 24))


    
    def show_details_frame():
        next_btn.place_forget()  
        main_frame.place_forget()
        details_frame.place(relx=0.5, rely=0.015, relwidth=0.9, relheight=0.89, anchor="n")  # Stretch to full height

    def back_to_main():
        details_frame.place_forget()
        main_frame.place(relx=0.5, rely=0.08, relwidth=0.9, anchor="n")
        next_btn.place(relx=0.5, rely=0.82, anchor="n", relwidth=0.9) 

        progress_bar.place_forget()
        status_label.lower()


    def finish_export():
        export_btn.configure(state="disabled")
        app.update()  # Force GUI refresh for disabled state

        metadata = {
            "instructor": instructor_entry.get().strip(),
            "course": course_entry.get().strip(),
            "semester": semester_entry.get().strip(),
            "notes": notes_entry.get().strip()
        }

        selected_files = [path for var, path in file_vars if var.get()]
        author = author_entry.get().strip()
        out_folder = output_folder_entry.get().strip() or "."
        file_name = output_file_entry.get().strip() or DEFAULT_FILE_NAME

        if not file_name.lower().endswith(".docx"):
            file_name += ".docx"

        output_path = os.path.join(out_folder, file_name)

        try:
            execution_outputs, visual_outputs = run_files_and_capture_output(selected_files)
            success = export_programs(author, selected_files, output_path, metadata, execution_outputs, visual_outputs)

            if success:
                show_confetti(success=True)
            else:
                show_confetti(success=False, message="Export Failed! Could not save the Word file.")
        except Exception as e:
            print(f"Export Failed: {e}")
            show_confetti(success=False, message="Export Failed! See terminal.")
        finally:
            export_btn.configure(state="normal")
            app.update()



    file_vars = []
    checkboxes = []

    # Title
    title_label = ctk.CTkLabel(app, text="AutoExport", font=("Segoe UI", 28, "bold"))
    title_label.place(relx=0.5, rely=0.015, anchor="n")

    # FRAME 1: MAIN INPUT
    main_frame = ctk.CTkFrame(app, corner_radius=12)
    main_frame.place(relx=0.5, rely=0.08, relwidth=0.9, anchor="n")

    def refresh_file_list(*args):
        for widget in scroll_frame.winfo_children():
            widget.destroy()
        file_vars.clear()
        checkboxes.clear()

        folders = prog_folder_entry.get().strip().split(";")
        SUPPORTED_EXTS = (".py", ".java", ".cpp", ".c", ".js", ".ts", ".html", ".css", ".go", ".rb", ".php", ".cs", ".sh", ".rs")

        files = []
        for folder in folders:
            folder = folder.strip()
            if os.path.isdir(folder):
                for root, _, files_in_dir in os.walk(folder):
                    for f in files_in_dir:
                        if f.endswith(SUPPORTED_EXTS):
                            path = os.path.join(root, f)
                            mod_time = os.path.getmtime(path)
                            size = os.path.getsize(path)
                            files.append((path, mod_time, size))

        # Sorting/filtering
        sort_option = sort_selector.get() if 'sort_selector' in globals() else "Date Modified"

        if sort_option == "Date Modified":
            files.sort(key=lambda x: x[1], reverse=True)
        elif sort_option == "File Size":
            files.sort(key=lambda x: x[2], reverse=True)
        elif sort_option.startswith("Language: "):
            lang = sort_option.split(": ")[1].lower()
            ext_map = {
                "python": ".py", "java": ".java", "c++": ".cpp", "c": ".c",
                "javascript": ".js", "typescript": ".ts", "html": ".html", "css": ".css",
                "go": ".go", "ruby": ".rb", "php": ".php", "c#": ".cs", "shell": ".sh", "rust": ".rs"
            }
            ext = ext_map.get(lang)
            if ext:
                files = [f for f in files if f[0].lower().endswith(ext)]

        search_term = search_entry.get().lower()
        found_any = False
        for path, mod_time, _ in files:
            filename = os.path.basename(path)
            if search_term in filename.lower():
                found_any = True
                var = ctk.BooleanVar(value=True)
                cb = ctk.CTkCheckBox(
                    scroll_frame,
                    text=f"{filename} ({datetime.datetime.fromtimestamp(mod_time).strftime('%d %b %Y')})",
                    variable=var
                )
                cb.pack(anchor="w", padx=10, pady=2)
                file_vars.append((var, path))
                checkboxes.append(cb)

        if not found_any:
            empty_label = ctk.CTkLabel(scroll_frame, text="No files found matching the criteria.", text_color="gray")
            empty_label.pack(pady=20)

        scroll_frame._parent_canvas.yview_moveto(0)


    def browse_folder(entry_widget):
        folder_path = filedialog.askdirectory()
        if folder_path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, folder_path)
            highlight_required(entry_widget, False)
            refresh_file_list()

    def toggle_select_all():
        all_selected = all(var.get() for var, _ in file_vars)
        for var, _ in file_vars:
            var.set(not all_selected)

    def validate_and_show_details():
        author = author_entry.get().strip()
        prog_folders = prog_folder_entry.get().strip()
        output_folder = output_folder_entry.get().strip()
        file_name = output_file_entry.get().strip()

        def is_valid_author(name):
            return len(name) >= 4 and all(c.isalpha() or c.isspace() for c in name)

        def is_valid_output_file(name):
            return name.lower().endswith(".docx") and len(name) > 1

        def is_valid_folder_paths(paths):
            return any(os.path.isdir(path.strip()) for path in paths.split(";"))

        def is_valid_output_folder(folder):
            return folder == "" or os.path.isdir(folder)

        valid = True

        if not is_valid_author(author):
            highlight_required(author_entry, True)
            valid = False
        else:
            highlight_required(author_entry, False)

        if not is_valid_folder_paths(prog_folders):
            highlight_required(prog_folder_entry, True)
            valid = False
        else:
            highlight_required(prog_folder_entry, False)

        if not is_valid_output_file(file_name):
            highlight_required(output_file_entry, True)
            valid = False
        else:
            highlight_required(output_file_entry, False)

        if not is_valid_output_folder(output_folder):
            highlight_required(output_folder_entry, True)
            valid = False
        else:
            highlight_required(output_folder_entry, False)

        if not valid:
            return

        save_config({"author": author, "folder": prog_folders, "outfolder": output_folder, "output": file_name})
        show_details_frame()
    
    author_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Shresth Dwivedi", height=32)
    prog_folder_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Programs or path/to/code", height=32)
    output_folder_entry = ctk.CTkEntry(main_frame, placeholder_text="Export folder (default: .)", height=32)
    output_file_entry = ctk.CTkEntry(main_frame, placeholder_text="e.g. Assignment.docx", height=32)

    attach_tooltip(author_entry, "Enter your full name (at least 4 letters)")
    attach_tooltip(prog_folder_entry, "Use ';' to separate multiple folders")
    attach_tooltip(output_file_entry, "Must end with .docx (e.g. Assignment.docx)")

    # Load saved values only if they exist
    if author := config.get("author"):
        author_entry.insert(0, author)

    if folder := config.get("folder"):
        prog_folder_entry.insert(0, folder)

    if outfolder := config.get("outfolder"):
        output_folder_entry.insert(0, outfolder)

    if output := config.get("output"):
        output_file_entry.insert(0, output)

    # Enable live validation
    bind_validation(author_entry, "author")
    bind_validation(prog_folder_entry, "prog_folder")
    bind_validation(output_folder_entry, "output_folder")
    bind_validation(output_file_entry, "output_file")
    
    labels = ["Author Name", "Programs Folder(s)", "Export Folder", "Output File Name (.docx)"]
    entries = [author_entry, prog_folder_entry, output_folder_entry, output_file_entry]
    for i, (label, entry) in enumerate(zip(labels, entries)):
        ctk.CTkLabel(main_frame, text=label).grid(row=i, column=0, sticky="w", padx=20, pady=(15 if i == 0 else 8))
        entry.grid(row=i, column=1, padx=20, pady=(15 if i == 0 else 8), sticky="ew")
        if "Folder" in label:
            ctk.CTkButton(main_frame, width=36, height=30, text="", image=folder_icon, command=lambda e=entry: browse_folder(e)).grid(row=i, column=2, padx=(0,20), sticky="ew")

    author_entry.grid(row=0, column=1, columnspan=2, padx=20, pady=(15 if i == 0 else 8), sticky="ew")
    output_file_entry.grid(row=3, column=1, columnspan=2, padx=20, pady=(15 if i == 0 else 8), sticky="ew")

    main_frame.grid_columnconfigure(1, weight=1)
    
    search_row = ctk.CTkFrame(main_frame, fg_color="transparent")
    search_row.grid(row=4, column=0, columnspan=3, padx=10, pady=(10, 5), sticky="ew")
    search_row.grid_columnconfigure(0, weight=1)  # Let search_entry expand

    search_entry = ctk.CTkEntry(search_row, placeholder_text="Search filename...", width=486)
    search_entry.grid(row=4, column=0, columnspan=2, padx=(10), pady=(10, 5), sticky="w")
    search_entry.bind("<KeyRelease>", refresh_file_list)

    toggle_btn = ctk.CTkButton(search_row, width=36, height=30, text="", image=checkmark_icon, command=toggle_select_all)
    toggle_btn.grid(row=4, column=2, padx=(0,10), pady=(10, 5), sticky="ew")

    sort_options = [
        "Date Modified", "File Size",
        "Language: Python", "Language: Java", "Language: C++", "Language: C",
        "Language: JavaScript", "Language: TypeScript", "Language: HTML", "Language: CSS",
        "Language: Go", "Language: Ruby", "Language: PHP", "Language: C#", "Language: Shell", "Language: Rust"
    ]
    sort_selector = ctk.CTkOptionMenu(
        master=search_row,
        values=sort_options,
        width=162,
        command=refresh_file_list
    )
    sort_selector.set("Date Modified")
    sort_selector.grid(row=4, column=1, columnspan=1, padx=21, pady=(10, 5), sticky="e")

    scroll_frame = ctk.CTkScrollableFrame(main_frame, height=200)
    scroll_frame.grid(row=5, column=0, columnspan=3, padx=20, pady=(5, 15), sticky="nsew")

    next_btn = ctk.CTkButton(app, text="Next", font=("Segoe UI", 16), fg_color="green", height=45, command=validate_and_show_details)
    next_btn.place(relx=0.5, rely=0.82, anchor="n", relwidth=0.9)

    progress_bar = ctk.CTkProgressBar(app, mode="determinate")
    progress_bar.place(relx=0.5, rely=0.89, anchor="center", relwidth=0.85)
    progress_bar.set(0)
    progress_bar.configure(determinate_speed=1)
    progress_bar.place_forget()

    status_label = ctk.CTkLabel(app, text="", font=("Segoe UI", 12), text_color="green")
    status_label.place(relx=0.5, rely=0.92, anchor="center")  
    status_label.lower() 

    # FRAME 2: DETAILS FRAME
    details_frame = ctk.CTkFrame(app, corner_radius=12)

    instructor_entry = ctk.CTkEntry(details_frame, placeholder_text="Instructor's Name")
    course_entry = ctk.CTkEntry(details_frame, placeholder_text="Course / Subject Name")
    semester_entry = ctk.CTkEntry(details_frame, placeholder_text="Semester or Session")
    notes_entry = ctk.CTkEntry(details_frame, placeholder_text="Additional Notes (optional)")

    meta_fields = [("Instructor", instructor_entry),
                   ("Course / Subject", course_entry),
                   ("Semester", semester_entry),
                   ("Notes", notes_entry)]

    for i, (label, entry) in enumerate(meta_fields):
        ctk.CTkLabel(details_frame, text=label).grid(row=i, column=0, sticky="w", padx=20, pady=(10 if i == 0 else 5))
        entry.grid(row=i, column=1, padx=20, pady=(10 if i == 0 else 5), sticky="ew")

    details_frame.grid_columnconfigure(1, weight=1)

    export_btn = ctk.CTkButton(details_frame, text="Export", font=("Segoe UI", 16), fg_color="green", height=45, command=lambda: threading.Thread(target=finish_export, daemon=True).start())
    export_btn.grid(row=6, column=0, columnspan=2, padx=20, pady=(10, 20), sticky="ew")

    back_btn = ctk.CTkButton(details_frame, text="Back", image=back_icon, command=back_to_main)
    back_btn.grid(row=7, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")

    def copy_upi_to_clipboard():
        app.clipboard_clear()
        app.clipboard_append("shresthdwivedi03@axl")
        app.update()

        copy_btn.configure(text=" Copied!", image=checkmark_icon, fg_color="green", hover_color="green")
        app.after(2000, lambda: copy_btn.configure(text=" Copy UPI", image=clipboard_icon, fg_color="#2b2b2b",hover_color="#144870"))

    donation_frame = ctk.CTkFrame(details_frame, fg_color="#333333", corner_radius=10)
    donation_frame.grid(row=10, column=0, columnspan=2, pady=(115, 10), padx=20, sticky="ew")

    support_label = ctk.CTkLabel(
        donation_frame,
        text="Support via UPI or follow me here:",
        font=("Segoe UI", 13, "italic"),
        #text_color="lightgreen",
        anchor="center"
    )
    support_label.pack(padx=10, pady=(10, 5))

    upi_id_label = ctk.CTkLabel(
        donation_frame,
        text="shresthdwivedi03@axl",
        font=("JetBrains Mono", 11),
        text_color="lightgreen"
    )
    upi_id_label.pack(pady=(0, 5))
    
    copy_btn = ctk.CTkButton(
        donation_frame,
        text=" Copy UPI",
        image=clipboard_icon,
        compound="left",
        fg_color="#2b2b2b",
        font=("Segoe UI", 12),
        command=copy_upi_to_clipboard
    )
    copy_btn.pack(pady=(0, 10))

    attach_tooltip(copy_btn, "Click to copy UPI ID")

    social_frame = ctk.CTkFrame(donation_frame, fg_color="transparent")
    social_frame.pack(pady=(5, 10))

    # GitHub icon button
    github_btn = ctk.CTkLabel(social_frame, text="", image=github_icon, cursor="hand2")
    github_btn.pack(side="left", padx=20, pady=(15,0))
    github_btn.bind("<Button-1>", lambda e: webbrowser.open("https://github.com/Shresth-Dwivedi"))

    # LinkedIn icon button
    linkedin_btn = ctk.CTkLabel(social_frame, text="", image=linkedin_icon, cursor="hand2")
    linkedin_btn.pack(side="left", padx=20, pady=(15,0))
    linkedin_btn.bind("<Button-1>", lambda e: webbrowser.open("https://www.linkedin.com/in/shresth-dwivedi/"))

    # Twitter icon button
    x_btn = ctk.CTkLabel(social_frame, text="", image=X_icon, cursor="hand2")
    x_btn.pack(side="left", padx=20, pady=(15,0))
    x_btn.bind("<Button-1>", lambda e: webbrowser.open("https://x.com/theDavyDee"))

    # Frame to center everything
    credit_frame = ctk.CTkFrame(app, fg_color="transparent")
    credit_frame.place(relx=0.5, rely=0.94, anchor="n")
    
    # Copyright label
    copyright_label = ctk.CTkLabel(
        credit_frame,
        text="© 2025 Shresth Dwivedi. All rights reserved.",
        text_color="white",
        font=("Segoe UI", 11)
    )
    copyright_label.pack()

    refresh_file_list()

    def show_update_popup():
        update_msg = check_for_update("1.5")
        if not update_msg:
            return

        popup = ctk.CTkButton(
            app,
            text=update_msg,
            image=bell_icon,
            text_color="white",
            fg_color="orange",
            hover_color="orange",
            corner_radius=8,
            font=("Segoe UI", 11),
            cursor="hand2"
        )

        popup.place(relx=0.5, rely=-0.1, anchor="n") 

        def open_release(event=None):
            webbrowser.open_new_tab("https://github.com/Shresth-Dwivedi/AutoExport/releases/latest")

        popup.bind("<Button-1>", open_release)

        def animate(step=0, reverse=False):
            if reverse:
                new_y = 0.07 - (step * 0.07)
            else:
                new_y = -0.1 + (step * 0.07)

            popup.place_configure(rely=new_y)

            if (not reverse and new_y < 0.07) or (reverse and new_y > -0.1):
                app.after(15, lambda: animate(step + 0.1, reverse))
            elif not reverse:
                app.after(10000, lambda: animate(0, reverse=True))

        animate()

    show_update_popup()

    app.mainloop()

if __name__ == "__main__":
    modern_gui()
